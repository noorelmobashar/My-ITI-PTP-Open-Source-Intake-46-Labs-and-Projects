"""
ingest.py — Document Ingestion Module
Loads PDFs, text, Word, JSON, CSV, and Markdown files from ./docs,
chunks them, and persists a FAISS index.
"""

import os
import gc
import json
import csv
import logging
from pathlib import Path

from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_ollama import OllamaEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.documents import Document

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------
DOCS_DIR = Path("./docs")
FAISS_INDEX_DIR = Path("./faiss_index")
CHUNK_SIZE = 1000
CHUNK_OVERLAP = 100
EMBEDDING_MODEL = "nomic-embed-text"
OLLAMA_BASE_URL = "http://localhost:11434"

# Supported file extensions for the UI and validation
SUPPORTED_EXTENSIONS = {
    ".pdf", ".txt", ".md", ".json", ".csv",
    ".docx", ".doc", ".log", ".xml", ".html", ".htm",
}

logging.basicConfig(level=logging.INFO, format="%(asctime)s  %(levelname)s  %(message)s")
logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Format-specific loaders
# ---------------------------------------------------------------------------

def _load_pdf(file_path: Path) -> list[Document]:
    loader = PyPDFLoader(str(file_path))
    return loader.load()


def _load_text(file_path: Path) -> list[Document]:
    """Load plain text, markdown, log, xml, html files."""
    loader = TextLoader(str(file_path), autodetect_encoding=True)
    return loader.load()


def _load_json(file_path: Path) -> list[Document]:
    """Load JSON — handles both objects and arrays, flattening to text."""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    documents = []
    if isinstance(data, list):
        for i, item in enumerate(data):
            text = json.dumps(item, indent=2, ensure_ascii=False) if isinstance(item, (dict, list)) else str(item)
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path.name, "item_index": i},
            ))
    elif isinstance(data, dict):
        text = json.dumps(data, indent=2, ensure_ascii=False)
        documents.append(Document(
            page_content=text,
            metadata={"source": file_path.name},
        ))
    else:
        documents.append(Document(
            page_content=str(data),
            metadata={"source": file_path.name},
        ))
    return documents


def _load_csv(file_path: Path) -> list[Document]:
    """Load CSV — each row becomes a document."""
    documents = []
    with open(file_path, "r", encoding="utf-8", newline="") as f:
        reader = csv.DictReader(f)
        for i, row in enumerate(reader):
            # Convert each row to a readable key: value format
            text = "\n".join(f"{k}: {v}" for k, v in row.items() if v)
            documents.append(Document(
                page_content=text,
                metadata={"source": file_path.name, "row": i},
            ))
    return documents


def _load_docx(file_path: Path) -> list[Document]:
    """Load Word .docx files using python-docx."""
    try:
        from docx import Document as DocxDocument
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return []

    doc = DocxDocument(str(file_path))
    full_text = "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    if not full_text.strip():
        return []
    return [Document(
        page_content=full_text,
        metadata={"source": file_path.name},
    )]


# Map extensions → loader functions
LOADER_MAP = {
    ".pdf":  _load_pdf,
    ".txt":  _load_text,
    ".md":   _load_text,
    ".log":  _load_text,
    ".xml":  _load_text,
    ".html": _load_text,
    ".htm":  _load_text,
    ".json": _load_json,
    ".csv":  _load_csv,
    ".docx": _load_docx,
    ".doc":  _load_docx,
}

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _load_documents(docs_dir: Path):
    """Discover and load all supported files under *docs_dir*."""
    documents = []

    if not docs_dir.exists():
        logger.warning("Docs directory %s does not exist – creating it.", docs_dir)
        docs_dir.mkdir(parents=True, exist_ok=True)
        return documents

    for file_path in sorted(docs_dir.iterdir()):
        if file_path.is_dir():
            continue
        suffix = file_path.suffix.lower()
        loader_fn = LOADER_MAP.get(suffix)
        if loader_fn is None:
            logger.debug("Skipping unsupported file: %s", file_path.name)
            continue
        try:
            logger.info("Loading %s …", file_path.name)
            docs = loader_fn(file_path)
            # Tag every chunk with the source filename for the UI
            for doc in docs:
                doc.metadata["source"] = file_path.name
            documents.extend(docs)
        except Exception as exc:
            logger.error("Failed to load %s: %s", file_path.name, exc)

    return documents


def _split_documents(documents):
    """Split documents into chunks."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        is_separator_regex=False,
    )
    chunks = splitter.split_documents(documents)
    logger.info("Split %d document(s) into %d chunks.", len(documents), len(chunks))
    return chunks


def _get_embeddings():
    """Return an OllamaEmbeddings instance."""
    return OllamaEmbeddings(
        model=EMBEDDING_MODEL,
        base_url=OLLAMA_BASE_URL,
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def ingest(docs_dir: Path = DOCS_DIR, index_dir: Path = FAISS_INDEX_DIR):
    """
    End-to-end ingestion pipeline.

    Returns
    -------
    dict  with keys ``num_documents``, ``num_chunks``, ``index_path``, ``sources``.
    """
    documents = _load_documents(docs_dir)
    if not documents:
        logger.warning("No documents found in %s.", docs_dir)
        return {
            "num_documents": 0,
            "num_chunks": 0,
            "index_path": str(index_dir),
            "sources": [],
        }

    chunks = _split_documents(documents)
    embeddings = _get_embeddings()

    logger.info("Building FAISS index (this may take a while) …")
    vectorstore = FAISS.from_documents(chunks, embeddings)

    index_dir.mkdir(parents=True, exist_ok=True)
    vectorstore.save_local(str(index_dir))
    logger.info("FAISS index saved to %s", index_dir)

    # Collect unique sources
    sources = sorted({doc.metadata.get("source", "unknown") for doc in documents})

    # Free GPU / CPU memory occupied by embedding batches
    del vectorstore
    gc.collect()

    return {
        "num_documents": len(documents),
        "num_chunks": len(chunks),
        "index_path": str(index_dir),
        "sources": sources,
    }


# ---------------------------------------------------------------------------
# Database Ingestion
# ---------------------------------------------------------------------------

DB_CONFIG = {
    "host": "localhost",
    "user": "noor",
    "password": "noor",
    "database": "iti",
    "port": 3306,
}


def load_from_database(db_config: dict = DB_CONFIG) -> list[Document]:
    """
    Connect to a MariaDB/MySQL database and extract:
      1. A schema overview document (all CREATE TABLE definitions)
      2. One document per table containing all rows as readable text

    This gives the LLM enough context to answer questions about
    both the database structure and the actual data.
    """
    import pymysql

    conn = pymysql.connect(
        host=db_config["host"],
        user=db_config["user"],
        password=db_config["password"],
        database=db_config["database"],
        port=db_config.get("port", 3306),
        cursorclass=pymysql.cursors.DictCursor,
    )
    documents = []
    db_name = db_config["database"]
    source_label = f"database:{db_name}"

    try:
        with conn.cursor() as cur:
            # ── Get all table names ──────────────────────────────
            cur.execute("SHOW TABLES")
            tables = [list(row.values())[0] for row in cur.fetchall()]
            logger.info("Database '%s' has %d tables: %s", db_name, len(tables), tables)

            # ── 1. Schema overview document ──────────────────────
            schema_parts = [f"Database: {db_name}\nTables: {', '.join(tables)}\n"]
            for table in tables:
                cur.execute(f"SHOW CREATE TABLE `{table}`")
                create_stmt = cur.fetchone()
                ddl = list(create_stmt.values())[1]
                schema_parts.append(f"\n--- Table: {table} ---\n{ddl}")

            documents.append(Document(
                page_content="\n".join(schema_parts),
                metadata={"source": source_label, "type": "schema"},
            ))

            # ── 2. One document per table with all row data ──────
            for table in tables:
                cur.execute(f"SELECT * FROM `{table}`")
                rows = cur.fetchall()
                if not rows:
                    documents.append(Document(
                        page_content=f"Table '{table}' in database '{db_name}' is empty (0 rows).",
                        metadata={"source": source_label, "type": "table_data", "table": table},
                    ))
                    continue

                columns = list(rows[0].keys())
                BATCH_SIZE = 5  # rows per document

                # Batch rows into groups, each document includes the table header
                for batch_start in range(0, len(rows), BATCH_SIZE):
                    batch = rows[batch_start:batch_start + BATCH_SIZE]
                    batch_end = batch_start + len(batch)

                    row_texts = []
                    for i, row in enumerate(batch, start=batch_start + 1):
                        fields = ", ".join(
                            f"{col}={row[col]}" for col in columns
                        )
                        row_texts.append(f"  Row {i}: {fields}")

                    content = (
                        f"Table: {table} (database: {db_name})\n"
                        f"Columns: {', '.join(columns)}\n"
                        f"Rows {batch_start + 1}-{batch_end} of {len(rows)}:\n"
                        + "\n".join(row_texts)
                    )
                    documents.append(Document(
                        page_content=content,
                        metadata={
                            "source": source_label,
                            "type": "table_data",
                            "table": table,
                            "rows": f"{batch_start + 1}-{batch_end}",
                        },
                    ))

            logger.info(
                "Extracted %d documents from database '%s' (%d tables).",
                len(documents), db_name, len(tables),
            )
    finally:
        conn.close()

    return documents


def ingest_database(db_config: dict = DB_CONFIG, index_dir: Path = FAISS_INDEX_DIR):
    """
    Pull all data from the database, chunk it, and merge into the
    existing FAISS index (or create a new one).
    """
    documents = load_from_database(db_config)
    if not documents:
        return {"num_documents": 0, "num_chunks": 0, "tables": []}

    # DB documents are already well-sized batches — skip the text splitter
    # to keep each table's data intact for accurate retrieval
    chunks = documents
    logger.info("Database produced %d documents (used as-is, no splitting).", len(chunks))
    embeddings = _get_embeddings()

    # If an index already exists, merge into it
    index_path = Path(index_dir)
    if index_path.exists() and (index_path / "index.faiss").exists():
        logger.info("Merging database chunks into existing FAISS index …")
        existing = FAISS.load_local(
            str(index_path), embeddings, allow_dangerous_deserialization=True
        )
        new_store = FAISS.from_documents(chunks, embeddings)
        existing.merge_from(new_store)
        existing.save_local(str(index_path))
        del existing, new_store
    else:
        logger.info("Creating new FAISS index from database data …")
        vectorstore = FAISS.from_documents(chunks, embeddings)
        index_path.mkdir(parents=True, exist_ok=True)
        vectorstore.save_local(str(index_path))
        del vectorstore

    gc.collect()

    tables = sorted({
        d.metadata.get("table", "schema")
        for d in documents
    })
    return {
        "num_documents": len(documents),
        "num_chunks": len(chunks),
        "tables": tables,
        "database": db_config["database"],
    }


# ---------------------------------------------------------------------------
# CLI entry-point
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    result = ingest()
    print("\n✅ Ingestion complete!")
    print(f"   Documents : {result['num_documents']}")
    print(f"   Chunks    : {result['num_chunks']}")
    print(f"   Index     : {result['index_path']}")
    print(f"   Sources   : {result['sources']}")
